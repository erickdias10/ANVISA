import re
from PyPDF2 import PdfReader
import unicodedata
from docx import Document
from docx.shared import Pt
import os
import streamlit as st
import subprocess
import sys

# ---------------------------
# Certifique-se de que o modelo SpaCy esteja baixado
# ---------------------------
def install_spacy_model():
    try:
        import spacy
        spacy.load("pt_core_news_sm")
    except OSError:
        subprocess.check_call([sys.executable, "-m", "spacy", "download", "pt_core_news_sm", "--user"])

install_spacy_model()
import spacy
nlp = spacy.load("pt_core_news_sm")

# ---------------------------
# Funções do SpaCy
# ---------------------------
def predict_with_spacy(text, entity_label):
    try:
        doc = nlp(text)
        entities = [ent.text for ent in doc.ents if ent.label_ == entity_label]
        return entities
    except Exception as e:
        print(f"Erro ao usar SpaCy para {entity_label}: {e}")
        return []

# ---------------------------
# Funções de Processamento de Texto
# ---------------------------
def normalize_text(text):
    if not isinstance(text, str):
        return text
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')
    text = re.sub(r"\s{2,}", " ", text)  # Remove múltiplos espaços
    return text.strip()

def corrigir_texto(texto):
    substituicoes = {
        'Ã©': 'é',
        'Ã§Ã£o': 'ção',
        'Ã³': 'ó',
        'Ã': 'à',
    }
    for errado, correto in substituicoes.items():
        texto = texto.replace(errado, correto)
    return texto

def extract_text_with_pypdf2(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        text = corrigir_texto(normalize_text(text))
        return text.strip()
    except Exception as e:
        print(f"Erro ao processar PDF {pdf_path}: {e}")
        return ''

# ---------------------------
# Funções de Extração de Dados
# ---------------------------
def extract_information_with_spacy(text):
    info = {
        "nome_autuado": predict_with_spacy(text, "PER") or None,
        "cnpj_cpf": predict_with_spacy(text, "MISC") or None,
        "socios_advogados": predict_with_spacy(text, "ORG") or [],
        "emails": predict_with_spacy(text, "EMAIL") or [],
    }
    return info

def extract_addresses_with_spacy(text):
    addresses = predict_with_spacy(text, "LOC")
    return [{"endereco": addr} for addr in addresses]

def extract_process_number(file_name):
    base_name = os.path.splitext(file_name)[0]  # Remove a extensão
    if base_name.startswith("SEI"):
        base_name = base_name[3:].strip()  # Remove "SEI"
    return base_name

# ---------------------------
# Função de Geração de Documento
# ---------------------------
def gerar_documento_docx(info, enderecos, numero_processo):
    try:
        output_directory = "output"
        os.makedirs(output_directory, exist_ok=True)

        output_path = os.path.join(output_directory, f"Notificacao_Processo_Nº_{numero_processo}.docx")

        doc = Document()
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        adicionar_paragrafo(doc, f"{info.get('nome_autuado', '[Nome não informado]')} – CNPJ/CPF: {info.get('cnpj_cpf', '[CNPJ/CPF não informado]')}")
        doc.add_paragraph("\n")

        for endereco in enderecos:
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            doc.add_paragraph("\n")

        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, "Informamos que foi proferido julgamento pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias no processo administrativo sancionador em referência.")

        doc.save(output_path)

        with open(output_path, "rb") as file:
            st.download_button(
                label="Baixar Documento",
                data=file,
                file_name=f"Notificacao_Processo_Nº_{numero_processo}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"Erro ao gerar o documento DOCX: {e}")

# ---------------------------
# Interface Streamlit
# ---------------------------
st.title("Sistema de Extração e Geração de Notificações")

uploaded_file = st.file_uploader("Envie um arquivo PDF", type="pdf")

if uploaded_file:
    try:
        file_name = uploaded_file.name
        numero_processo = extract_process_number(file_name)
        text = extract_text_with_pypdf2(uploaded_file)
        if text:
            st.success(f"Texto extraído com sucesso! Número do processo: {numero_processo}")

            info = extract_information_with_spacy(text) or {}
            addresses = extract_addresses_with_spacy(text) or []

            if st.button("Gerar Documento"):
                gerar_documento_docx(info, addresses, numero_processo)
    except Exception as e:
        st.error(f"Ocorreu um erro: {e}")
