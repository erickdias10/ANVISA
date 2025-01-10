import streamlit as st
import logging
import asyncio
import time
import os
import unicodedata
import re
import spacy
import difflib
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from io import BytesIO

# Adicionando novas importações para OCR
from pdf2image import convert_from_path
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter

# Configuração de logs
logging.basicConfig(level=logging.INFO)

# Configurar o caminho do Tesseract (ajuste conforme necessário)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Para Windows
# pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'  # Para Linux

# Configurar a política de loop de eventos para Windows
if os.name == 'nt':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

# Constantes de elementos
LOGIN_URL = "https://sei.anvisa.gov.br/sip/login.php?sigla_orgao_sistema=ANVISA&sigla_sistema=SEI"

def create_browser_context():
    """
    Configura e retorna uma instância do Playwright BrowserContext e Page,
    utilizando um diretório de dados de usuário persistente para manter a sessão.
    """
    download_dir = os.path.join(os.getcwd(), "downloads")
    os.makedirs(download_dir, exist_ok=True)

    # Diretório para armazenar dados de usuário persistentes
    user_data_dir = os.path.join(os.getcwd(), "user_data")
    os.makedirs(user_data_dir, exist_ok=True)

    playwright = sync_playwright().start()
    context = playwright.chromium.launch_persistent_context(
        user_data_dir=user_data_dir,
        headless=False,  # Alterado para False para depuração
        accept_downloads=True,  # Permite downloads automáticos
        downloads_path=download_dir  # Define o diretório de downloads
    )
    page = context.new_page()
    return playwright, context, page

def wait_for_element(page, selector, timeout=20000):
    """
    Aguarda até que um elemento esteja presente no DOM.
    """
    try:
        logging.info(f"Aguardando elemento: {selector}")
        element = page.wait_for_selector(selector, timeout=timeout)
        if element:
            logging.info(f"Elemento {selector} encontrado.")
            return element
    except PlaywrightTimeoutError:
        logging.error(f"Erro ao localizar o elemento: {selector}")
        raise Exception(f"Elemento {selector} não encontrado na página.")
    return None

def handle_download(download, download_dir):
    """
    Manipula o download, salvando-o no diretório de downloads especificado.
    """
    os.makedirs(download_dir, exist_ok=True)
    download_path = os.path.join(download_dir, download.suggested_filename)
    download.save_as(download_path)
    logging.info(f"Download salvo em: {download_path}")
    return download_path

def handle_alert(page):
    """
    Captura e trata alertas inesperados sem recarregar a página.
    """
    try:
        dialog = page.expect_event("dialog", timeout=5000)
        if dialog:
            alert_text = dialog.message
            logging.warning(f"Alerta inesperado encontrado: {alert_text}")
            dialog.accept()
            return alert_text
    except PlaywrightTimeoutError:
        logging.info("Nenhum alerta encontrado.")
        return None

def login(page, username, password):
    """
    Realiza o login no sistema SEI.
    """
    logging.info("Acessando a página de login.")
    page.goto(LOGIN_URL)

    # Aguarda o campo de usuário aparecer e preenche
    user_field = wait_for_element(page, "#txtUsuario")
    if user_field:
        user_field.fill(username)
        logging.info("Campo de usuário preenchido.")
    else:
        raise Exception("Campo de usuário não encontrado.")

    # Aguarda o campo de senha aparecer e preenche
    password_field = wait_for_element(page, "#pwdSenha")
    if password_field:
        password_field.fill(password)
        logging.info("Campo de senha preenchido.")
    else:
        raise Exception("Campo de senha não encontrado.")

    # Aguarda o botão de login aparecer e clica
    login_button = wait_for_element(page, "#sbmAcessar")
    if login_button:
        login_button.click()
        logging.info("Botão de login clicado.")
    else:
        raise Exception("Botão de login não encontrado.")

    # Aguarda a página principal carregar após login
    try:
        page.wait_for_load_state("networkidle", timeout=20000)  # Aumentei o timeout
        logging.info("Login realizado com sucesso.")

        # Captura uma screenshot após o login para depuração
        screenshot_path = os.path.join(os.getcwd(), "login_success.png")
        page.screenshot(path=screenshot_path)
        logging.info(f"Screenshot após login salva em: {screenshot_path}")
    except PlaywrightTimeoutError:
        # Captura uma screenshot em caso de falha no login
        screenshot_path = os.path.join(os.getcwd(), "login_failure.png")
        page.screenshot(path=screenshot_path)
        logging.error("Tempo esgotado aguardando a página principal carregar após login.")
        raise Exception("Login pode não ter sido realizado com sucesso. Verifique a screenshot para mais detalhes.")

def access_process(page, process_number):
    """
    Acessa um processo pelo número no sistema SEI.
    """
    try:
        search_field = wait_for_element(page, "#txtPesquisaRapida", timeout=40000)  # Aumentado para 40 segundos
        search_field.fill(process_number)
        logging.info(f"Preenchido o número do processo: {process_number}")
        search_field.press("Enter")
        logging.info("Processo acessado com sucesso.")
        time.sleep(5)  # Aumentado para 5 segundos

        # Captura uma screenshot após acessar o processo
        screenshot_path = os.path.join(os.getcwd(), "access_process_success.png")
        page.screenshot(path=screenshot_path)
        logging.info(f"Screenshot após acessar o processo salva em: {screenshot_path}")

    except Exception as e:
        # Captura uma screenshot para depuração
        screenshot_path = os.path.join(os.getcwd(), "access_process_failure.png")
        page.screenshot(path=screenshot_path)
        logging.error(f"Erro ao acessar o processo: {e}")
        raise Exception(f"Erro ao acessar o processo: {e}. Screenshot salva em {screenshot_path}")

# Defina seus identificadores e XPaths
IFRAME_VISUALIZACAO_ID = "ifrVisualizacao"
BUTTON_XPATH_GERAR_PDF = '//*[@id="divArvoreAcoes"]/a[7]/img'  # XPath para o botão 'Gerar PDF'
BUTTON_XPATH_DOWNLOAD_OPTION = '//*[@id="divInfraBarraComandosSuperior"]/button[1]'  # XPath para a opção de download

def generate_and_download_pdf(page, download_dir):
    """
    Gera e baixa o PDF do processo no iframe correspondente.
    :param page: Instância do Playwright Page.
    :param download_dir: Diretório para salvar o download.
    """
    try:
        # Espera pelo iframe e obtém o handle
        logging.info(f"Esperando pelo iframe com ID {IFRAME_VISUALIZACAO_ID}")
        iframe_element = page.wait_for_selector(f'iframe#{IFRAME_VISUALIZACAO_ID}', timeout=10000)
        if not iframe_element:
            raise Exception(f"Iframe com ID {IFRAME_VISUALIZACAO_ID} não encontrado.")
    
        # Acessa o contexto do iframe
        iframe = iframe_element.content_frame()
        if not iframe:
            raise Exception("Não foi possível acessar o conteúdo do iframe.")
    
        logging.info("Iframe encontrado e acessado com sucesso.")
    
        # Espera pelo botão de gerar PDF dentro do iframe
        logging.info(f"Esperando pelo botão de gerar PDF com XPath {BUTTON_XPATH_GERAR_PDF}")
        gerar_pdf_button = iframe.wait_for_selector(f'xpath={BUTTON_XPATH_GERAR_PDF}', timeout=10000)
        if not gerar_pdf_button:
            raise Exception("Botão para gerar PDF não encontrado.")
    
        # Clicar no botão 'Gerar PDF' para abrir as opções de download
        logging.info("Clicando no botão 'Gerar Arquivo PDF do Processo'")
        gerar_pdf_button.click()
        logging.info("Clique no botão 'Gerar Arquivo PDF do Processo' realizado.")
    
        # Opcional: esperar que as opções de download apareçam
        time.sleep(2)  # Aguarda um breve momento para que as opções de download apareçam
    
        # Espera pelo botão de opção de download dentro do iframe
        logging.info(f"Esperando pelo botão de opção de download com XPath {BUTTON_XPATH_DOWNLOAD_OPTION}")
        download_option_button = iframe.wait_for_selector(f'xpath={BUTTON_XPATH_DOWNLOAD_OPTION}', timeout=10000)
        if not download_option_button:
            raise Exception("Botão de opção de download não encontrado.")
    
        # Clicar no botão de opção de download e capturar o download
        logging.info("Clicando no botão de opção de download.")
        with page.expect_download(timeout=60000) as download_info_option:
            download_option_button.click()
        download_option = download_info_option.value
        download_option_path = handle_download(download_option, download_dir)
        logging.info("Clique no botão de opção de download realizado.")
    
        # Retorna o caminho do download principal
        return download_option_path
    
    except PlaywrightTimeoutError as e:
        logging.error(f"Timeout ao gerar o PDF: {e}")
        raise Exception("Timeout ao gerar o PDF do processo.")
    except Exception as e:
        logging.error(f"Erro ao gerar o PDF: {e}")
        raise Exception("Erro ao gerar o PDF do processo.")
    finally:
        # Opcional: espera adicional se necessário
        time.sleep(5)

def process_notification(username, password, process_number):
    """
    Orquestra o processo de login, acesso ao processo e geração/baixa do PDF.
    """
    playwright, context, page = create_browser_context()
    download_dir = os.path.join(os.getcwd(), "downloads")
    try:
        # Passo 1: Login
        login(page, username, password)

        # Passo 2: Acessa o processo
        access_process(page, process_number)

        # Passo 3: Gera e baixa o PDF
        try:
            download_path = generate_and_download_pdf(page, download_dir)  # Função consolidada
            logging.info(f"PDF gerado e salvo em: {download_path}")
        except Exception as e:
            logging.error(f"Erro ao gerar o PDF: {e}")
            raise Exception("Erro durante o processo de geração do PDF.")

        logging.info("PDF gerado com sucesso.")

        return download_path  # Retorna o caminho do PDF baixado
    except Exception as e:
        logging.error(f"Erro durante o processamento: {e}")
        raise e
    finally:
        # Fechar o navegador
        context.close()
        playwright.stop()

def normalize_text(text):
    if not isinstance(text, str):
        return text
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')
    text = re.sub(r"\s{2,}", " ", text)  # Remove múltiplos espaços
    return text.strip()

def corrigir_texto(texto):
    substituicoes = {
        'Ã©': 'é',
        'Ã§Ã£o': 'ção',
        'Ã³': 'ó',
        'Ã': 'à',
        'â€“': '–',
        'â€”': '—',
        'Ãº': 'ú',
        'Ãª': 'ê',
        'Ã£o': 'ão',
        'â€œ': '"',
        'â€': '"',
        'Ã¡': 'á',
        'Ã¢': 'â',
        'Ã­': 'í',
        'Ã´': 'ô',
        'Ã§': 'ç',
        # Adicione mais substituições conforme necessário
    }
    for errado, correto in substituicoes.items():
        texto = texto.replace(errado, correto)
    return texto

def extract_information_spacy(text):
    """
    Extrai informações do texto utilizando spaCy.
    """
    doc = nlp(text)

    info = {
        "nome_autuado": None,
        "cpf": None,
        "cnpj": None,
        "socios_advogados": [],
        "emails": [],
    }

    for ent in doc.ents:
        if ent.label_ in ["PER", "ORG"]:  # Pessoa ou Organização
            if not info["nome_autuado"]:
                info["nome_autuado"] = ent.text.strip()
        elif ent.label_ == "EMAIL":
            info["emails"].append(ent.text.strip())

    # Usar regex para complementar a extração de CNPJ e CPF
    cnpj_pattern = r"CNPJ:\s*([\d./-]{18})"
    cpf_pattern = r"CPF:\s*([\d./-]{14})"

    cnpj_match = re.search(cnpj_pattern, text)
    cpf_match = re.search(cpf_pattern, text)

    if cnpj_match:
        cnpj = cnpj_match.group(1)
        info["cnpj"] = format_cnpj(cnpj)
    if cpf_match:
        cpf = cpf_match.group(1)
        info["cpf"] = format_cpf(cpf)

    # Sócios ou Advogados mencionados
    socios_adv_pattern = r"(?:Sócio|Advogado|Responsável|Representante Legal):\s*([\w\s]+)"
    info["socios_advogados"] = re.findall(socios_adv_pattern, text) or []

    return info

def format_cnpj(cnpj):
    """
    Formata o CNPJ no padrão XX.XXX.XXX/XXXX-XX
    """
    digits = re.sub(r'\D', '', cnpj)
    if len(digits) != 14:
        return cnpj  # Retorna como está se não tiver 14 dígitos
    return f"{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:14]}"

def format_cpf(cpf):
    """
    Formata o CPF no padrão XXX.XXX.XXX-XX
    """
    digits = re.sub(r'\D', '', cpf)
    if len(digits) != 11:
        return cpf  # Retorna como está se não tiver 11 dígitos
    return f"{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:11]}"

def extract_process_number(file_name):
    base_name = os.path.splitext(file_name)[0]
    if base_name.startswith("SEI"):
        base_name = base_name[3:].strip()
    # Formatar no padrão XXXXX.XXXXXX/XXXX-XX
    digits = re.sub(r'\D', '', base_name)
    if len(digits) != 15:
        return base_name  # Retorna como está se não tiver 15 dígitos
    return f"{digits[:5]}.{digits[5:11]}/{digits[11:15]}-{digits[15-1:]}"

def extract_text_with_pypdf2(pdf_path):
    """
    Extrai texto de um PDF usando PyPDF2. Se não encontrar texto, tenta usar OCR.
    """
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text

        if text.strip():
            logging.info("Texto extraído com sucesso usando PyPDF2.")
            text = corrigir_texto(normalize_text(text))
            return text.strip()
        else:
            logging.info("Nenhum texto encontrado com PyPDF2. Tentando OCR.")
            # Usar OCR
            return extract_text_with_ocr(pdf_path)
    except Exception as e:
        st.error(f"Erro ao processar PDF {pdf_path}: {e}")
        return ''

def extract_text_with_ocr(pdf_path):
    """
    Extrai texto de um PDF baseado em imagem usando OCR.
    """
    try:
        # Converter PDF para imagens com maior DPI para melhor qualidade
        logging.info("Convertendo PDF para imagens para OCR.")
        pages = convert_from_path(pdf_path, dpi=300, fmt='jpeg')  # Usando JPEG para melhor compressão

        text = ""
        for page_number, page in enumerate(pages, start=1):
            logging.info(f"Processando página {page_number} com OCR.")
            # Pré-processamento da imagem para melhorar a precisão do OCR
            gray = page.convert('L')  # Converter para escala de cinza
            enhancer = ImageEnhance.Contrast(gray)
            gray = enhancer.enhance(2.0)
            threshold = gray.point(lambda x: 0 if x < 128 else 255, '1')  # Aplicar threshold binário
            threshold = threshold.filter(ImageFilter.MedianFilter())

            custom_config = r'--oem 3 --psm 6'  # OEM 3: LSTM + Legacy, PSM 6: Assume uma única unidade de texto uniforme
            page_text = pytesseract.image_to_string(threshold, lang='por', config=custom_config)
            text += page_text + "\n"

        if text.strip():
            logging.info("Texto extraído com sucesso usando OCR.")
            text = corrigir_texto(normalize_text(text))
            return text.strip()
        else:
            logging.warning("Nenhum texto encontrado mesmo após OCR.")
            return ''
    except Exception as e:
        st.error(f"Erro durante o OCR do PDF {pdf_path}: {e}")
        return ''

def extract_addresses_spacy(text):
    """
    Extrai endereços do texto utilizando spaCy e complementa com regex, garantindo a unicidade e completude.
    - Exclui endereços com Endereço: ou CEP: vazios ou nulos.
    - Seleciona o endereço com mais caracteres para cada grupo baseado na normalização.
    - Preenche os outros campos (cidade, bairro, estado) da melhor forma possível.
    """
    doc = nlp(text)

    addresses = []
    seen_addresses = {}

    # Padrões regex específicos
    endereco_pattern = r"(?:Endereço|End|Endereco):\s*([\w\s.,ºª-]+)"
    cidade_pattern = r"Cidade:\s*([\w\s]+(?: DE [\w\s]+)?)"
    bairro_pattern = r"Bairro:\s*([\w\s]+)"
    estado_pattern = r"Estado:\s*([A-Z]{2})"
    cep_pattern = r"CEP:\s*(\d{2}\.\d{3}-\d{3}|\d{5}-\d{3})"

    endereco_matches = re.findall(endereco_pattern, text)
    cidade_matches = re.findall(cidade_pattern, text)
    bairro_matches = re.findall(bairro_pattern, text)
    estado_matches = re.findall(estado_pattern, text)
    cep_matches = re.findall(cep_pattern, text)

    max_len = max(len(endereco_matches), len(cidade_matches), len(bairro_matches), len(estado_matches), len(cep_matches))

    for i in range(max_len):
        endereco = endereco_matches[i].strip() if i < len(endereco_matches) else "[Não informado]"
        cidade = cidade_matches[i].strip() if i < len(cidade_matches) else "[Não informado]"
        bairro = bairro_matches[i].strip() if i < len(bairro_matches) else "[Não informado]"
        estado = estado_matches[i].strip() if i < len(estado_matches) else "[Não informado]"
        cep = cep_matches[i].strip() if i < len(cep_matches) else "[Não informado]"

        # Verificar se 'Endereço' e 'CEP' não estão vazios ou com placeholders
        if (endereco and endereco != "[Não informado]") and (cep and cep != "[Não informado]"):
            # Normalizar o endereço para deduplicação
            normalized_endereco = normalize_address(endereco)

            # Só se adiciona se o endereço for mais completo
            if normalized_endereco in seen_addresses:
                existing = seen_addresses[normalized_endereco]
                # Seleciona o endereço com mais caracteres
                if len(endereco) > len(existing["endereco"]):
                    seen_addresses[normalized_endereco]["endereco"] = endereco

                # Preenche os campos restantes se estiverem faltando
                if existing["cidade"] == "[Não informado]" and cidade != "[Não informado]":
                    seen_addresses[normalized_endereco]["cidade"] = cidade
                if existing["bairro"] == "[Não informado]" and bairro != "[Não informado]":
                    seen_addresses[normalized_endereco]["bairro"] = bairro
                if existing["estado"] == "[Não informado]" and estado != "[Não informado]":
                    seen_addresses[normalized_endereco]["estado"] = estado
                if existing["cep"] == "[Não informado]" and cep != "[Não informado]":
                    seen_addresses[normalized_endereco]["cep"] = cep
            else:
                seen_addresses[normalized_endereco] = {
                    "endereco": endereco,
                    "cidade": cidade,
                    "bairro": bairro,
                    "estado": estado,
                    "cep": cep
                }

    # Converter o dicionário para uma lista de endereços
    for addr in seen_addresses.values():
        addresses.append({
            "endereco": addr["endereco"],
            "cidade": addr["cidade"],
            "bairro": addr["bairro"],
            "estado": addr["estado"],
            "cep": addr["cep"]
        })

    return addresses

def normalize_address(address):
    """
    Normaliza o endereço removendo pontuação, convertendo para minúsculas e removendo espaços extras.
    """
    address = unicodedata.normalize('NFKD', address).encode('ASCII', 'ignore').decode('utf-8')
    address = re.sub(r'[^\w\s]', '', address)  # Remove pontuação
    address = re.sub(r'\s+', ' ', address)  # Remove múltiplos espaços
    return address.lower().strip()

def adicionar_paragrafo(doc, texto="", negrito=False, tamanho=12):
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run(texto)
    run.bold = negrito
    run.font.size = Pt(tamanho)
    return paragrafo

def _gerar_modelo_1(doc, info, enderecos, numero_processo, email_selecionado):
    try:
        # Adiciona o cabeçalho do documento
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        nome_autuado = info.get('nome_autuado', '[Nome não informado]')
        cnpj = info.get('cnpj', '')
        cpf = info.get('cpf', '')
        if cnpj:
            identificador = f"CNPJ: {cnpj}"
        elif cpf:
            identificador = f"CPF: {cpf}"
        else:
            identificador = "CNPJ/CPF: [Não informado]"
        adicionar_paragrafo(doc, f"{nome_autuado} – {identificador}")
        doc.add_paragraph("\n")

        for idx, endereco in enumerate(enderecos, start=1):
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Cidade: {endereco.get('cidade', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Bairro: {endereco.get('bairro', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Estado: {endereco.get('estado', '[Não informado]')}")
            adicionar_paragrafo(doc, f"CEP: {endereco.get('cep', '[Não informado]')}")
            doc.add_paragraph("\n")

        adicionar_paragrafo(doc, 
            "Assunto: Decisão de 1ª instância proferida pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias.", 
            negrito=True
        )
        adicionar_paragrafo(doc, 
            f"Referência: Processo Administrativo Sancionador nº: {numero_processo} ", 
            negrito=True
        )
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, 
            "Informamos que foi proferido julgamento pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias no processo administrativo sancionador em referência, conforme decisão em anexo."
        )
        doc.add_paragraph("\n")

        adicionar_paragrafo(doc, "O QUE FAZER SE A DECISÃO TIVER APLICADO MULTA?", negrito=True)
        adicionar_paragrafo(doc, 
            "Sendo aplicada a penalidade de multa, esta notificação estará acompanhada de boleto bancário, que deverá ser pago até o vencimento."
        )
        adicionar_paragrafo(doc, 
            "O valor da multa poderá ser pago com 20% de desconto caso seja efetuado em até 20 dias contados de seu recebimento. "
            "Incorrerá em ilegalidade o usufruto do desconto em data posterior ao prazo referido, mesmo que a data impressa no boleto permita pagamento, "
            "sendo a diferença cobrada posteriormente pela Gerência de Gestão de Arrecadação (GEGAR). "
            "O pagamento da multa implica em desistência tácita do recurso, conforme art. 21 da Lei nº 6.437/1977."
        )
        adicionar_paragrafo(doc, 
            "O não pagamento do boleto sem que haja interposição de recurso, acarretará, sucessivamente: "
            "i) a inscrição do devedor no Cadastro Informativo de Crédito não Quitado do Setor Público Federal (CADIN); "
            "ii) a inscrição do débito em dívida ativa da União; iii) o ajuizamento de ação de execução fiscal contra o devedor; "
            "e iv) a comunicação aos cartórios de registros de imóveis, dos devedores inscritos em dívida ativa ou execução fiscal."
        )
        adicionar_paragrafo(doc, 
            "Esclarecemos que o valor da multa foi atualizado pela taxa Selic acumulada nos termos do art. 37-A da Lei 10.522/2002 "
            "e no art. 5º do Decreto-Lei 1.736/79."
        )
        doc.add_paragraph("\n")

        adicionar_paragrafo(doc, "COMO FAÇO PARA INTERPOR RECURSO DA DECISÃO?", negrito=True)
        adicionar_paragrafo(doc, 
            "Havendo interesse na interposição de recurso administrativo, este poderá ser interposto no prazo de 20 dias contados do recebimento desta notificação, "
            "conforme disposto no art. 9º da RDC nº 266/2019."
        )
        adicionar_paragrafo(doc, 
            "O protocolo do recurso deverá ser feito exclusivamente, por meio de peticionamento intercorrente no processo indicado no campo assunto desta notificação, "
            "pelo Sistema Eletrônico de Informações (SEI). Para tanto, é necessário, primeiramente, fazer o cadastro como usuário externo SEI-Anvisa. "
            "Acesse o portal da Anvisa https://www.gov.br/anvisa/pt-br > Sistemas > SEI > Acesso para Usuários Externos (SEI) e siga as orientações. "
            "Para maiores informações, consulte o Manual do Usuário Externo Sei-Anvisa, que está disponível em https://www.gov.br/anvisa/pt-br/sistemas/sei."
        )
        doc.add_paragraph("\n")

        adicionar_paragrafo(doc, "QUAIS DOCUMENTOS DEVEM ACOMPANHAR O RECURSO?", negrito=True)
        adicionar_paragrafo(doc, "a) Autuado pessoa jurídica:")
        adicionar_paragrafo(doc, "1. Contrato ou estatuto social da empresa, com a última alteração;")
        adicionar_paragrafo(doc, 
            "2. Procuração e documento de identificação do outorgado (advogado ou representante), caso constituído para atuar no processo. "
            "Somente serão aceitas procurações e substabelecimentos assinados eletronicamente, com certificação digital no padrão da "
            "Infraestrutura de Chaves Públicas Brasileira (ICP-Brasil) ou pelo assinador Gov.br."
        )
        adicionar_paragrafo(doc, 
            "3. Ata de eleição da atual diretoria quando a procuração estiver assinada por diretor que não conste como sócio da empresa;"
        )
        adicionar_paragrafo(doc, 
            "4. No caso de contestação sobre o porte da empresa considerado para a dosimetria da pena de multa: comprovação do porte econômico "
            "referente ao ano em que foi proferida a decisão (documentos previstos no art. 50 da RDC nº 222/2006)."
        )
        adicionar_paragrafo(doc, "b) Autuado pessoa física:")
        adicionar_paragrafo(doc, "1. Documento de identificação do autuado;")
        adicionar_paragrafo(doc, 
            "2. Procuração e documento de identificação do outorgado (advogado ou representante), caso constituído para atuar no processo."
        )
        adicionar_paragrafo(doc, f"\nInformações de contato: {email_selecionado}")
    
    except Exception as e:
        st.error(f"Erro ao gerar o documento no modelo 1: {e}")

def _gerar_modelo_2(doc, info, enderecos, numero_processo, motivo_revisao, data_decisao, data_recebimento_notificacao, data_extincao=None, email_selecionado=None):
    try:
        # Adiciona o cabeçalho do documento
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        nome_autuado = info.get('nome_autuado', '[Nome não informado]')
        cnpj = info.get('cnpj', '')
        cpf = info.get('cpf', '')
        if cnpj:
            identificador = f"CNPJ: {cnpj}"
        elif cpf:
            identificador = f"CPF: {cpf}"
        else:
            identificador = "CNPJ/CPF: [Não informado]"
        adicionar_paragrafo(doc, f"{nome_autuado} – {identificador}")
        doc.add_paragraph("\n")

        for idx, endereco in enumerate(enderecos, start=1):
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Cidade: {endereco.get('cidade', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Bairro: {endereco.get('bairro', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Estado: {endereco.get('estado', '[Não informado]')}")
            adicionar_paragrafo(doc, f"CEP: {endereco.get('cep', '[Não informado]')}")
            doc.add_paragraph("\n")

        adicionar_paragrafo(doc, 
            "Assunto: Decisão de 1ª instância proferida pela Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias.", 
            negrito=True
        )
        adicionar_paragrafo(doc, 
            f"Referência: Processo Administrativo Sancionador nº: {numero_processo} ", 
            negrito=True
        )
        doc.add_paragraph("\n")
        
        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        
        # Texto Adaptado
        adicionar_paragrafo(doc, 
            f"Informamos que a Decisão em 1ª instância proferida pela Gerência-Geral de Portos, Aeroportos, Fronteiras e Recintos Alfandegados ou Coordenação de Atuação Administrativa e Julgamento das Infrações Sanitárias, em {data_decisao.strftime('%d/%m/%Y')}, no processo administrativo sancionador em referência, foi revisada ou retratada no âmbito da Anvisa pelos motivos expostos abaixo."
        )
        doc.add_paragraph("\n")
        
        if motivo_revisao == "insuficiencia_provas":
            adicionar_paragrafo(doc, 
                "Foi constatado que não há comprovação suficiente nos autos do processo para afirmar que a recorrente cometeu a infração objeto da autuação em questão."
            )
        elif motivo_revisao == "prescricao":
            adicionar_paragrafo(doc, 
                f"Foi observado que da decisão condenatória recorrível proferida em {data_decisao.strftime('%d/%m/%Y')} até o ato seguinte capaz de interromper a prescrição (ex: notificação da decisão em {data_recebimento_notificacao.strftime('%d/%m/%Y')}) passaram-se mais de cinco anos sem que houvesse entre eles outro ato capaz de interromper o curso prescricional (documento que declarou a prescrição. Ex: NOTA n. 00014/2020/EI-M-ANVIS/ENAC/PGF/AGU)."
            )
        elif motivo_revisao == "extincao_empresa":
            if not data_extincao:
                raise ValueError("A data de extinção da empresa deve ser fornecida para o motivo 'extincao_empresa'.")
            adicionar_paragrafo(doc, 
                f"Foi constatado, ao longo dos procedimentos de cobrança administrativa, que a empresa em questão havia sido 'EXTINTA' na data de {data_extincao.strftime('%d/%m/%Y')}, conforme Certidão Simplificada e documento de Distrato Social fornecido pelo órgão de registro comercial - [Nome do Órgão]."
            )
        else:
            # Para outros motivos, você pode adaptar conforme necessário
            adicionar_paragrafo(doc, 
                "Foi constatado que há razões adicionais para a revisão/retratação da decisão, conforme detalhado nos documentos anexos."
            )
        
        doc.add_paragraph("\n")
        adicionar_paragrafo(doc, 
            "Dessa forma, a decisão condenatória perdeu seus efeitos e o processo será arquivado."
        )
        doc.add_paragraph("\n")
        
        # Seção "Como Obter Cópia do Processo"
        adicionar_paragrafo(doc, "COMO OBTER CÓPIA DO PROCESSO?", negrito=True)
        adicionar_paragrafo(doc, 
            "Informações e pedidos de cópias devem ser solicitados exclusivamente pelos Canais de Atendimento da Anvisa (https://www.gov.br/anvisa/pt-br/canais_atendimento) ou pelo Serviço de Atendimento ao Cidadão (https://www.gov.br/anvisa/pt-br/acessoainformacao/sic)."
        )
        adicionar_paragrafo(doc, 
            "Os pedidos de cópia de processo devem informar o número do processo e a finalidade da cópia."
        )
        adicionar_paragrafo(doc, 
            "A cópia integral dos autos somente será concedida para o interessado direto no processo, ou seu representante devidamente constituído, cuja condição deve ser comprovada mediante a apresentação dos seguintes documentos:"
        )
        adicionar_paragrafo(doc, "1. Documento de identificação do autuado (se pessoa física) ou outorgado;")
        adicionar_paragrafo(doc, 
            "2. Procuração e documento de identificação do outorgado (advogado ou representante), caso seja ele o requerente. "
            "Somente serão aceitas procurações e substabelecimento assinados eletronicamente, com certificação digital no padrão da Infraestrutura de Chaves Públicas Brasileira (ICP-Brasil) ou pelo assinador Gov.br."
        )
        adicionar_paragrafo(doc, 
            "3. Contrato ou estatuto social da empresa, com a última alteração (se pessoa jurídica);"
        )
        adicionar_paragrafo(doc, 
            "4. Ata de eleição da atual diretoria quando a procuração estiver assinada por diretor que não conste como sócio da empresa (se pessoa jurídica);"
        )
        adicionar_paragrafo(doc, 
            "A ausência de quaisquer dos documentos acima ensejará o indeferimento sumário do pedido."
        )
        adicionar_paragrafo(doc, 
            "Terceiros não interessados diretamente no processo estão dispensados de apresentar documentação e terão acesso somente às cópias dos seguintes documentos: Auto de Infração, Manifestação da área autuante e Decisão."
        )
        adicionar_paragrafo(doc, f"\nInformações de contato: {email_selecionado}")

    except Exception as e:
        st.error(f"Erro ao gerar o documento no modelo 2: {e}")

def _gerar_modelo_3(doc, info, enderecos, numero_processo, usuario_nome, usuario_email, orgao_registro_comercial, email_selecionado):
    try:
        # Adiciona uma quebra de linha no início
        doc.add_paragraph("\n")
        
        # Adiciona o cabeçalho do documento
        adicionar_paragrafo(doc, "Ao(a) Senhor(a):")
        nome_autuado = info.get('nome_autuado', '[Nome não informado]')
        cnpj = info.get('cnpj', '')
        cpf = info.get('cpf', '')
        if cnpj:
            identificador = f"CNPJ: {cnpj}"
        elif cpf:
            identificador = f"CPF: {cpf}"
        else:
            identificador = "CNPJ/CPF: [Não informado]"
        adicionar_paragrafo(doc, f"{nome_autuado} – {identificador}")
        doc.add_paragraph("\n")

        # Adiciona os endereços
        for idx, endereco in enumerate(enderecos, start=1):
            adicionar_paragrafo(doc, f"Endereço: {endereco.get('endereco', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Cidade: {endereco.get('cidade', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Bairro: {endereco.get('bairro', '[Não informado]')}")
            adicionar_paragrafo(doc, f"Estado: {endereco.get('estado', '[Não informado]')}")
            adicionar_paragrafo(doc, f"CEP: {endereco.get('cep', '[Não informado]')}")
            doc.add_paragraph("\n")
        
        # Adiciona o assunto e referência em negrito
        adicionar_paragrafo(doc, "Assunto: Decisão proferida pela Diretoria Colegiada", negrito=True)
        adicionar_paragrafo(doc, f"Referência: Processo Administrativo Sancionador nº {numero_processo}", negrito=True)
        doc.add_paragraph("\n")
        
        # Saudação
        adicionar_paragrafo(doc, "Prezado(a) Senhor(a),")
        doc.add_paragraph("\n")
        
        # Corpo do documento
        adicionar_paragrafo(doc, 
            "Informamos que foi proferido julgamento da Diretoria Colegiada no processo administrativo sancionador em referência, conforme decisão em anexo, contra a qual não cabe recurso."
        )
        adicionar_paragrafo(doc, "\n")
        
        adicionar_paragrafo(doc, 
            "Em sendo mantida a penalidade de multa, esta notificação estará acompanhada de boleto bancário. Exceto para a decisão, cujo recurso tenha sido considerado intempestivo, um vez que o boleto será encaminhado pela Gerência de Gestão de Arrecadação – GEGAR."
        )
        adicionar_paragrafo(doc, "\n")
        
        adicionar_paragrafo(doc, 
            "O não pagamento do boleto, caso devido, acarretará, sucessivamente: i) a inscrição do devedor no Cadastro Informativo de Crédito não Quitado do Setor Público Federal (CADIN); ii) a inscrição do débito em dívida ativa da União; iii) o ajuizamento de ação de execução fiscal contra o devedor; e iv) a comunicação aos cartórios de registros de imóveis, dos devedores inscritos em dívida ativa ou execução fiscal."
        )
        adicionar_paragrafo(doc, "\n")
        
        adicionar_paragrafo(doc, 
            "Esclarecemos que, em caso de penalidade de multa, seu valor foi atualizado pela taxa Selic acumulada nos termos do art. 37-A da Lei 10.522/2002 e no art. 5º do Decreto-Lei 1.736/79."
        )
        adicionar_paragrafo(doc, "\n")
        
        # Seção "Informações e pedidos de cópias"
        adicionar_paragrafo(doc, 
            "Informações e pedidos de cópias podem ser solicitados pelos Canais de Atendimento da Anvisa (webchat, formulário eletrônico ou telefone 0800 642 9782), responsáveis por atender a esse tipo de demanda de forma centralizada. Os pedidos de cópia de PAS devem vir acompanhados dos documentos abaixo, sob pena de não serem atendidos:"
        )
        adicionar_paragrafo(doc, "\n")
        
        # Lista de documentos necessários
        documentos = [
            "Cópia autenticada da procuração/substabelecimento com firma reconhecida e poderes específicos para tal;",
            "Cópia do CPF e do RG do outorgado e do requerente, caso sejam pessoas distintas; e",
            "Cópia autenticada do contrato social/estatuto social, com a última alteração."
        ]
        
        for doc_item in documentos:
            adicionar_paragrafo(doc, f"- {doc_item}")
        adicionar_paragrafo(doc, "\n")
        
        adicionar_paragrafo(doc, 
            f"Por fim, esclarecemos que foi concedido aos autos por meio do Sistema Eletrônico de Informações (SEI), por 180 (cento e oitenta) dias, ao usuário: {usuario_nome} ({usuario_email})."
        )
        adicionar_paragrafo(doc, "\n")
        
        # Informações de contato
        adicionar_paragrafo(doc, f"\nInformações de contato: {email_selecionado}")
        
        # Encerramento
        adicionar_paragrafo(doc, "Atenciosamente,")
        adicionar_paragrafo(doc, "\n")
        adicionar_paragrafo(doc, f"{usuario_nome}")
        
    except Exception as e:
        st.error(f"Erro ao gerar o documento no modelo 3: {e}")

def extract_all_emails(emails):
    """
    Remove duplicatas e retorna uma lista de emails únicos.
    """
    return list(set(emails))

def main():
    st.title("Gerador de Notificações SEI-Anvisa")

    st.sidebar.header("Informações de Login")
    username = st.sidebar.text_input("Usuário")
    password = st.sidebar.text_input("Senha", type="password")

    st.header("Processo Administrativo")
    process_number_input = st.text_input("Número do Processo")

    # Botão para gerar notificação
    gerar_notificacao = st.button("Gerar Notificação")

    if gerar_notificacao:
        if not username or not password or not process_number_input:
            st.error("Por favor, preencha todos os campos.")
        else:
            with st.spinner("Processando..."):
                try:
                    download_path = process_notification(username, password, process_number_input)
                    st.success("PDF gerado com sucesso!")

                    if download_path:
                        # Extrair o nome do arquivo PDF
                        pdf_file_name = os.path.basename(download_path)
                        st.info(f"Arquivo PDF baixado: {pdf_file_name}")

                        # Extrair o número do processo e formatar
                        numero_processo = extract_process_number(pdf_file_name)

                        # Extrair texto do PDF
                        text = extract_text_with_pypdf2(download_path)

                        if text:
                            st.success("Texto extraído com sucesso!")
                            info = extract_information_spacy(text)
                            addresses = extract_addresses_spacy(text)
                            emails = extract_all_emails(info.get('emails', []))

                            # Exibir informações extraídas
                            st.subheader("Informações Extraídas")
                            st.write(f"**Arquivo PDF:** {pdf_file_name}")  # Informar de qual arquivo a informação foi extraída
                            st.write(f"**Nome Autuado:** {info.get('nome_autuado', 'Não informado')}")
                            if info.get('cnpj'):
                                st.write(f"**CNPJ:** {info.get('cnpj')}")
                            elif info.get('cpf'):
                                st.write(f"**CPF:** {info.get('cpf')}")
                            st.write(f"**Emails:** {', '.join(emails) if emails else 'Não informado'}")
                            st.write(f"**Sócios/Advogados:** {', '.join(info.get('socios_advogados', []))}")

                            st.subheader("Endereços Encontrados")
                            for idx, end in enumerate(addresses, start=1):
                                st.write(f"**Endereço {idx}:**")
                                st.write(f"  - Endereço: {end['endereco']}")
                                st.write(f"  - Cidade: {end['cidade']}")
                                st.write(f"  - Bairro: {end['bairro']}")
                                st.write(f"  - Estado: {end['estado']}")
                                st.write(f"  - CEP: {end['cep']}")

                            # Permitir ao usuário editar os endereços
                            st.subheader("Editar Endereços")
                            edited_addresses = []
                            for idx, end in enumerate(addresses, start=1):
                                st.write(f"**Endereço {idx}:**")
                                endereco = st.text_input(f"Endereço {idx}", value=end['endereco'])
                                cidade = st.text_input(f"Cidade {idx}", value=end['cidade'])
                                bairro = st.text_input(f"Bairro {idx}", value=end['bairro'])
                                estado = st.text_input(f"Estado {idx}", value=end['estado'])
                                cep = st.text_input(f"CEP {idx}", value=end['cep'])
                                edited_addresses.append({
                                    "endereco": endereco,
                                    "cidade": cidade,
                                    "bairro": bairro,
                                    "estado": estado,
                                    "cep": cep
                                })

                            # Permitir ao usuário selecionar qual email utilizar
                            st.subheader("Selecionar Email para Utilizar no Processo")
                            if emails:
                                email_selecionado = st.selectbox("Selecione o email desejado:", emails)
                            else:
                                email_selecionado = "[Não informado]"

                            # Armazenar informações no session_state para uso posterior
                            st.session_state['info'] = info
                            st.session_state['enderecos'] = edited_addresses
                            st.session_state['numero_processo'] = numero_processo
                            st.session_state['pdf_file'] = pdf_file_name
                            st.session_state['email_selecionado'] = email_selecionado

                except Exception as ex:
                    st.error(f"Ocorreu um erro: {ex}")

    # Verificar se a notificação foi gerada e armazenada no session_state
    if ('info' in st.session_state and 
        'enderecos' in st.session_state and 
        'numero_processo' in st.session_state and
        'pdf_file' in st.session_state):
        
        st.subheader("Escolha o Modelo do Documento")
        modelo = st.selectbox("Selecione o modelo desejado:", [
            "ID 3791 - Notificação de decisões em 1ª instância - SEI", 
            "ID 2782 - Notificação de decisões revisadas/retratadas", 
            "ID 2703 - Notificação de decisão da DICOL"
        ])

        gerar_doc = st.button("Gerar Documento Word")

        if gerar_doc:
            try:
                doc = Document()
                info = st.session_state['info']
                edited_addresses = st.session_state['enderecos']
                numero_processo = st.session_state['numero_processo']
                pdf_file_name = st.session_state['pdf_file']  # Obter o nome do arquivo PDF
                email_selecionado = st.session_state['email_selecionado']

                # Solicitar informações adicionais conforme o modelo selecionado
                if modelo == "ID 3791 - Notificação de decisões em 1ª instância - SEI":
                    _gerar_modelo_1(doc, info, edited_addresses, numero_processo, email_selecionado)
                elif modelo == "ID 2782 - Notificação de decisões revisadas/retratadas":
                    # Solicitar informações adicionais necessárias para o modelo 2
                    motivo_revisao = st.selectbox("Motivo da Revisão:", ["insuficiencia_provas", "prescricao", "extincao_empresa", "outros"])
                    data_decisao = st.date_input("Data da Decisão:")
                    data_recebimento_notificacao = st.date_input("Data de Recebimento da Notificação:")
                    
                    data_extincao = None
                    if motivo_revisao == "extincao_empresa":
                        data_extincao = st.date_input("Data de Extinção da Empresa:")

                    _gerar_modelo_2(
                        doc, 
                        info, 
                        edited_addresses, 
                        numero_processo, 
                        motivo_revisao, 
                        data_decisao, 
                        data_recebimento_notificacao, 
                        data_extincao,
                        email_selecionado
                    )
                elif modelo == "ID 2703 - Notificação de decisão da DICOL":
                    # Solicitar informações adicionais necessárias para o modelo 3
                    usuario_nome = st.text_input("Nome do Usuário:")
                    usuario_email = st.text_input("Email do Usuário:")
                    orgao_registro_comercial = st.text_input("Órgão de Registro Comercial:")
                    _gerar_modelo_3(
                        doc, 
                        info, 
                        edited_addresses, 
                        numero_processo, 
                        usuario_nome, 
                        usuario_email, 
                        orgao_registro_comercial,
                        email_selecionado
                    )

                # Adicionar uma nota sobre a origem das informações no documento
                adicionar_paragrafo(doc, f"\nInformações extraídas do arquivo PDF: {pdf_file_name}", negrito=False, tamanho=10)

                # Salvar documento em buffer
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # Nome do arquivo
                modelo_id = re.findall(r'ID\s(\d+)', modelo)[0] if re.findall(r'ID\s(\d+)', modelo) else "unknown"
                output_filename = f"Notificacao_Processo_Nº_{numero_processo}_modelo_{modelo_id}.docx"

                st.download_button(
                    label="Baixar Documento",
                    data=buffer,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("Documento gerado e pronto para download!")

            except Exception as ex:
                st.error(f"Ocorreu um erro ao gerar o documento: {ex}")

if __name__ == '__main__':
    # Carregar o modelo spaCy para português
    try:
        nlp = spacy.load("pt_core_news_lg")
    except OSError:
        st.info("Modelo 'pt_core_news_lg' não encontrado. Instalando...")
        os.system("python -m spacy download pt_core_news_lg")
        nlp = spacy.load("pt_core_news_lg")

    main()
